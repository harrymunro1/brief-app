# from langchain_community.document_loaders import Docx2txtLoader
from langchain.prompts import PromptTemplate
from docx import Document
# import re
#import os
from langchain_openai import OpenAIEmbeddings
from langchain_openai import ChatOpenAI
import streamlit as st
from dotenv import load_dotenv
# import base64
from io import BytesIO


# Load environment variables from .env file
load_dotenv()

# Function to create a docx in memory
def create_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer





## Create the top line 
st.write("# Tool for interpreting client briefs")
st.write("""This tool is desinged to help CS understand which AA projects would be best suited to tackle thier clients needs.""")
st.write("## Inputs")

# Get api key
#api_key = os.getenv("OPENAI_API_KEY")
api_key = st.secrets.get("OPENAI_API_KEY")

# Get project master and brief
uploaded_file = st.file_uploader("Upload **project master**, stored here `W:\Commercial\Expert Solutions\Brief Interpretation\Master`",
                                 type=["docx"])
uploaded_file_proj = st.file_uploader("Upload a **brief**, stored here `W:\Commercial\Expert Solutions\Brief Interpretation\Briefs`",
                                 type=["docx"])


# Allow for the output file to be named
output_file = st.text_input("What would you like the output to be saved as?")



## Add a part that allows the user to input the clients budget 
st.write("##### Client **Budget** Input")
use_budget = st.checkbox("Add budget constraints?")

if use_budget:
    # Budget input fields

    # Budget input fields
    min_budget = st.number_input("Minimum Budget (£000's)", min_value=0, value=0, step=10)
    max_budget = st.number_input("Maximum Budget (£000's)", min_value=min_budget, value=100, step=10)
    
    # Display the result
    st.markdown(f"**The client's budget is between £{min_budget*1000:,} and £{max_budget*1000:,}.**")


    ## Need to change the budget numbers to 1000's
    min_budget = min_budget*1000
    max_budget = max_budget*1000




## Button to execute response
pressed = st.button("Press to load response and save output")



# Create embeddings and store in vector database
embeddings = OpenAIEmbeddings(openai_api_key=api_key)

## Main response function
if pressed:  
    
    # brief_loader = Docx2txtLoader(uploaded_file)
    # brief_doc = brief_loader.load()
    # brief_text = brief_doc[0].page_content
    
    # Load the Word document
    doc = Document(uploaded_file)
   
   # Extract text
    brief_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Read in project file
    doc_proj = Document(uploaded_file_proj)
   
   # Extract text
    projects_context = "\n".join([para.text for para in doc_proj.paragraphs])
    
    # projects_context = documents[0].page_content
    if use_budget:
        # Create prompt for the LLM
        prompt = PromptTemplate(
            input_variables=["brief", "projects","max_budget","min_budget"],
            template="""
            You are an expert consultant tasked with matching client needs to appropriate projects.
            
            CLIENT BRIEF:
            {brief}
            
            POTENTIAL RELEVANT PROJECTS:
            {projects}
            
            Based on the client brief and the project descriptions provided, which project would be the best match for this client? 
            Make sure your recommendation is deifinitly within the clients budget of between {min_budget} and {max_budget}.
            Explain your reasoning clearly, highlighting how the recommended project addresses the client's specific needs, occasionally mentioning specific numbers this could help with from the client brief.
            If multiple projects could work together, suggest that combination approach however if none are suited then responed accordingly, and certainly dont suggest projects that will noy help the client.
            
            Can you make sure to also include the email and contact name as well as the budget required for each of the projects you have suggested.
            """
        )
        # Get LLM recommendation
        llm = ChatOpenAI(
                    temperature=0.3,
                    model_name="gpt-4o",
                    # model_name="gpt-4",
                    #model_name="gpt-4-turbo",
                    max_tokens=1000,
                     openai_api_key=api_key
                )
        
            # Create the chain
        chain = prompt | llm
        
        # Use `.invoke()` instead of `.run()`
        response = chain.invoke({
            "brief": brief_text,
            "projects": projects_context,
            "max_budget": max_budget,
            "min_budget": min_budget
            })
            
    else: 
        # Create prompt for the LLM
        prompt = PromptTemplate(
            input_variables=["brief", "projects"],
            template=
            """
            You are an expert consultant tasked with matching client needs to appropriate projects.
            
            CLIENT BRIEF:
            {brief}
            
            POTENTIAL RELEVANT PROJECTS:
            {projects}
            
            Based on the client brief and the project descriptions provided, which project would be the best match for this client? 
            Explain your reasoning clearly, highlighting how the recommended project addresses the client's specific needs, occasionally mentioning specific numbers this could help with from the client brief.
            If multiple projects could work together, suggest that combination approach however if none are suited then responed accordingly, and certainly dont suggest projects that will noy help the client.
            
            Can you make sure to also include the email and contact name as well as the budget required for each of the projects you have suggested.
            """
        )
        # Get LLM recommendation
        llm = ChatOpenAI(
                    temperature=0.3,
                    model_name="gpt-4o",
                    #model_name="o1-preview",
                    # model_name="gpt-4",
                    #model_name="gpt-4-turbo",
                    max_tokens=1000,
                     openai_api_key=api_key
                )
        
            # Create the chain
        chain = prompt | llm
        
        # Use `.invoke()` instead of `.run()`
        response = chain.invoke({
            "brief": brief_text,
            "projects": projects_context

    })
    
    
    
    ## Write out reponse
    st.text(f"{response.content}")
    
    # if save_down:
        # Generate the docx file and provide a download button
    docx_file = create_docx(response.content)
    st.download_button(
        label="📄 Download as Word (.docx)",
        data=docx_file,
        file_name=output_file,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


    # ## Email prompt
    # # grab all the emails from the response
    # emails = list(set(re.findall(r'\b[\w\.-]+@Kantar\.com\b', response.content)),
    #               set(re.findall(r'\b[\w\.-]+@kantar\.com\b', response.content)))
    # st.write(emails)





